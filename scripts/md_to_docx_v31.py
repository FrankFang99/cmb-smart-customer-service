# -*- coding: utf-8 -*-
"""
评测评分标准 v3.1 → docx 生成器
输入: D:\Learning\AI\面试\AI智能客服\docs\评测评分标准_v3.1.md
输出: 同目录 评测评分标准_v3.1.docx
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"D:\Learning\AI\面试\AI智能客服\docs\评测评分标准_v3.1.md")
DST = Path(r"D:\Learning\AI\面试\AI智能客服\docs\评测评分标准_v3.1.docx")

# 中文字体（统一）
CN_FONT = "微软雅黑"
EN_FONT = "Calibri"

def set_cn_font(run, size_pt=10.5, bold=False, color=None):
    run.font.name = EN_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), CN_FONT)
    rFonts.set(qn('w:ascii'), EN_FONT)
    rFonts.set(qn('w:hAnsi'), EN_FONT)

def add_heading(doc, text, level):
    """level: 1/2/3"""
    sizes = {1: 16, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_cn_font(run, size_pt=sizes[level], bold=True,
                color=(31, 73, 125) if level == 1 else (0, 0, 0))
    return p

def add_paragraph(doc, text, indent_cm=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    set_cn_font(run, size_pt=10.5)
    return p

def add_blockquote(doc, text):
    """引用块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # 浅灰背景
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(text)
    set_cn_font(run, size_pt=10, color=(89, 89, 89))
    return p

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BFBFBF')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + 0.5 * level)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    set_cn_font(run, size_pt=10.5)
    return p

def add_numbered(doc, text, num):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    r1 = p.add_run(f"{num}. ")
    set_cn_font(r1, size_pt=10.5, bold=True)
    r2 = p.add_run(text)
    set_cn_font(r2, size_pt=10.5)
    return p

def add_table(doc, rows, has_header=True):
    """rows: list[list[str]]"""
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = 'Light Grid Accent 1'
    table.autofit = True
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.2
            run = p.add_run(cell_text)
            if has_header and i == 0:
                set_cn_font(run, size_pt=10, bold=True, color=(255, 255, 255))
                # 表头底色
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '4472C4')
                tcPr.append(shd)
            else:
                set_cn_font(run, size_pt=9.5)
    return table

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F8F8F8')
    pPr.append(shd)
    # 左边框
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '12')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), '4472C4')
    pBdr.append(left)
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rPr.append(rFonts)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

def parse_md_table(lines, start):
    """从 start 行开始解析 markdown 表格，返回 (rows, next_idx)"""
    rows = []
    i = start
    while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
        line = lines[i].strip()
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip('|').split('|')]
        rows.append(cells)
        i += 1
    return rows, i

def clean_text(s):
    """去掉 markdown 加粗标记等"""
    return s.replace('**', '')

def main():
    doc = Document()
    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # 默认样式
    style = doc.styles['Normal']
    style.font.name = EN_FONT
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), CN_FONT)

    text = SRC.read_text(encoding='utf-8')
    lines = text.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行
        if not stripped:
            i += 1
            continue

        # 水平分割线
        if stripped == '---':
            add_hr(doc)
            i += 1
            continue

        # 标题
        if stripped.startswith('# '):
            add_heading(doc, clean_text(stripped[2:]), 1)
            i += 1
            continue
        if stripped.startswith('## '):
            add_heading(doc, clean_text(stripped[3:]), 2)
            i += 1
            continue
        if stripped.startswith('### '):
            add_heading(doc, clean_text(stripped[4:]), 3)
            i += 1
            continue
        if stripped.startswith('#### '):
            add_heading(doc, clean_text(stripped[5:]), 3)
            i += 1
            continue

        # 引用块
        if stripped.startswith('> '):
            add_blockquote(doc, clean_text(stripped[2:]))
            i += 1
            continue

        # 表格
        if stripped.startswith('|'):
            rows, i = parse_md_table(lines, i)
            if rows:
                add_table(doc, rows, has_header=True)
            continue

        # 代码块
        if stripped.startswith('```'):
            lang = stripped[3:].strip() or "text"
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            add_code(doc, '\n'.join(code_lines))
            continue

        # 列表项
        if re.match(r'^\s*[-*+]\s+', line):
            indent = len(line) - len(line.lstrip())
            level = 0 if indent < 2 else 1
            content = re.sub(r'^\s*[-*+]\s+', '', line)
            add_bullet(doc, clean_text(content), level=level)
            i += 1
            continue

        # 有序列表
        m = re.match(r'^\s*\d+\.\s+(.*)', line)
        if m:
            num_match = re.match(r'^\s*(\d+)\.', line)
            num = num_match.group(1)
            add_numbered(doc, clean_text(m.group(1)), num)
            i += 1
            continue

        # 普通段落
        # 收集连续非空行（段落）
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(('#', '|', '>', '---', '```', '-', '*', '+')) \
                and not re.match(r'^\s*\d+\.\s+', lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        add_paragraph(doc, clean_text(' '.join(para_lines)))

    doc.save(str(DST))
    print(f"OK: {DST}")
    print(f"Size: {DST.stat().st_size} bytes")

if __name__ == "__main__":
    main()
